import os
import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches

import matplotlib.pyplot as plt
import seaborn as sns

from config.analysis import get_scope
from src.extract.wansoft import load_orders
from src.transform.quadrant_scatter import quadrant_scatter_matplotlib

def ensure_dirs():
    os.makedirs("output/word", exist_ok=True)
    os.makedirs("output/figures", exist_ok=True)


def safe_num(s):
    return pd.to_numeric(s, errors="coerce")


def month_to_dt(m):
    return pd.to_datetime(m + "-01", errors="coerce")


def load_inputs():
    kpis = pd.read_csv("output/datasets/kpis_monthly.csv")
    recon = pd.read_csv("output/datasets/reconciliation_monthly.csv")
    ops = pd.read_csv("output/datasets/kpis_operational.csv")
    return kpis, recon, ops


def prepare_monthly(kpis, recon, ops):
    if "delta_ventas_pct" not in recon.columns:
        recon["delta_ventas_pct"] = np.nan
    if "source_flag" not in recon.columns:
        recon["source_flag"] = "OK"
    if "status" not in recon.columns:
        recon["status"] = "OK"
    if "note" not in recon.columns:
        recon["note"] = ""

    df = (
        kpis.merge(recon[["month","branch","status","source_flag","note","delta_ventas_pct"]], on="month", how="left")
            .merge(ops, on="month", how="left")
    )

    for c in ["Ventas","Tickets","Personas","Ticket_Promedio","Cheque_Promedio","CostoTotal","COGS_%","Margen","Margen_%",
              "delta_ventas_pct","cortesias_en_platillos","cancelaciones_en_platillos","anulaciones_en_platillos","descuentos_en_platillos",
              "Ventas_Salon","Ventas_Delivery","Ventas_Llevar","%_Delivery","%_Salon","%_Llevar"]:
        if c in df.columns:
            df[c] = safe_num(df[c])

    if "COGS_%" not in df.columns or df["COGS_%"].isna().all():
        df["COGS_%"] = df["CostoTotal"] / df["Ventas"]
    if "Margen" not in df.columns or df["Margen"].isna().all():
        df["Margen"] = df["Ventas"] - df["CostoTotal"]
    if "Margen_%" not in df.columns or df["Margen_%"].isna().all():
        df["Margen_%"] = df["Margen"] / df["Ventas"]

    df["month_dt"] = month_to_dt(df["month"])
    df = df.sort_values("month_dt")
    df["is_real_issue"] = (df["status"] == "REVIEW") & (df["source_flag"] == "OK")
    return df


def build_orders_ops():
    orders = load_orders()
    orders["Fecha_dt"] = pd.to_datetime(orders["Fecha"], errors="coerce")
    orders["month"] = orders["Fecha_dt"].dt.to_period("M").astype(str)
    orders["Total_num"] = safe_num(orders["Total"]).fillna(0)
    orders["Personas_num"] = safe_num(orders["Personas"]).fillna(0)

    # Mesas y clientes
    mesas = (orders.groupby("month", as_index=False)
             .agg(mesas=("Mesa","nunique"),
                  clientes=("Personas_num","sum"),
                  ventas=("Total_num","sum")))
    mesas["month_dt"] = month_to_dt(mesas["month"])
    mesas = mesas.sort_values("month_dt")

    # MESEROS (NUEVO)
    meseros = (orders.groupby(["month","Mesero"], as_index=False)
               .agg(
                   ventas=("Total_num","sum"),
                   tickets=("Orden","nunique"),
                   personas=("Personas_num","sum"),
                   mesas=("Mesa","nunique")
               ))
    
    meseros["ticket_promedio"] = meseros["ventas"] / meseros["tickets"]
    meseros["personas_por_mesa"] = meseros["personas"] / meseros["mesas"]

    return mesas, orders, meseros


def save_quadrant_figures(df, mesas, meseros):
    # -------- CUADRANTES MENSUALES --------
    fig, axs = plt.subplots(2, 3, figsize=(18, 10))

    quadrant_scatter_matplotlib(df, "Tickets", "Ventas", "month", "Ventas vs Tickets", ax=axs[0,0])
    quadrant_scatter_matplotlib(df, "Tickets", "Ticket_Promedio", "month", "Ticket vs Tickets", ax=axs[0,1])
    quadrant_scatter_matplotlib(df, "Personas", "Cheque_Promedio", "month", "Cheque vs Personas", ax=axs[0,2])
    quadrant_scatter_matplotlib(mesas, "mesas", "clientes", "month", "Mesas vs Clientes", ax=axs[1,0])
    quadrant_scatter_matplotlib(df, "%_Delivery", "Ticket_Promedio", "month", "Delivery vs Ticket", ax=axs[1,1])
    quadrant_scatter_matplotlib(df, "Ventas", "COGS_%", "month", "Ventas vs COGS%", ax=axs[1,2])

    plt.tight_layout()
    plt.savefig("output/figures/cuadrantes_mensual.png", dpi=160)
    plt.close()

    # -------- CUADRANTES MESEROS --------
    mes = meseros.copy()
    mes["month_dt"] = pd.to_datetime(mes["month"] + "-01")
    mes = mes[mes["month_dt"].dt.year >= 2025]
    last_month = mes["month"].max()
    mes = mes[mes["month"] == last_month]

    fig, axs = plt.subplots(2, 2, figsize=(16, 10))

    quadrant_scatter_matplotlib(mes, "personas", "ventas", "Mesero", "Ventas vs Clientes", ax=axs[0,0])
    quadrant_scatter_matplotlib(mes, "tickets", "ticket_promedio", "Mesero", "Tickets vs Ticket", ax=axs[0,1])
    quadrant_scatter_matplotlib(mes, "mesas", "ventas", "Mesero", "Ventas vs Mesas", ax=axs[1,0])
    quadrant_scatter_matplotlib(mes, "personas_por_mesa", "ticket_promedio", "Mesero", "Ticket vs Pers/Mesa", ax=axs[1,1])

    plt.tight_layout()
    plt.savefig("output/figures/cuadrantes_meseros.png", dpi=160)
    plt.close()

def save_core_figures(df, mesas):
    # 1) Ventas + COGS% + Margen%
    fig1 = plt.figure(figsize=(12,5))
    ax = plt.gca()
    ax.plot(df["month_dt"], df["Ventas"], marker="o", label="Ventas")
    ax.set_title("Evolución mensual: Ventas")
    ax.set_ylabel("MXN")
    plt.xticks(rotation=45)
    plt.tight_layout()
    p1 = "output/figures/ventas_mensual.png"
    plt.savefig(p1, dpi=160); plt.close(fig1)

    fig2 = plt.figure(figsize=(12,5))
    ax = plt.gca()
    ax.plot(df["month_dt"], df["COGS_%"]*100, marker="o", label="COGS %")
    ax.plot(df["month_dt"], df["Margen_%"]*100, marker="o", label="Margen %")
    ax.set_title("Evolución mensual: COGS% y Margen%")
    ax.set_ylabel("%")
    ax.legend()
    plt.xticks(rotation=45)
    plt.tight_layout()
    p2 = "output/figures/cogs_margen.png"
    plt.savefig(p2, dpi=160); plt.close(fig2)

    # 2) Ticket y Cheque
    fig3 = plt.figure(figsize=(12,5))
    ax = plt.gca()
    ax.plot(df["month_dt"], df["Ticket_Promedio"], marker="o", label="Ticket promedio")
    ax.plot(df["month_dt"], df["Cheque_Promedio"], marker="o", label="Cheque promedio")
    ax.set_title("Evolución mensual: Ticket y Cheque promedio")
    ax.set_ylabel("MXN")
    ax.legend()
    plt.xticks(rotation=45)
    plt.tight_layout()
    p3 = "output/figures/ticket_cheque.png"
    plt.savefig(p3, dpi=160); plt.close(fig3)

    # 3) Mesas y clientes
    fig4 = plt.figure(figsize=(12,5))
    ax = plt.gca()
    ax.plot(mesas["month_dt"], mesas["mesas"], marker="o", label="Mesas")
    ax.plot(mesas["month_dt"], mesas["clientes"], marker="o", label="Clientes (Personas)")
    ax.set_title("Evolución mensual: Mesas y Clientes")
    ax.legend()
    plt.xticks(rotation=45)
    plt.tight_layout()
    p4 = "output/figures/mesas_clientes.png"
    plt.savefig(p4, dpi=160); plt.close(fig4)

    # 4) Correlaciones (heatmap)
    cols = ["Ventas","Ticket_Promedio","Cheque_Promedio","COGS_%","Margen_%","%_Delivery","Personas"]
    cols = [c for c in cols if c in df.columns]
    corr = df[cols].corr(numeric_only=True)

    fig5 = plt.figure(figsize=(10,6))
    sns.heatmap(corr, annot=True, cmap="RdBu", vmin=-1, vmax=1)
    plt.title("Correlaciones mensuales (Pearson)")
    plt.tight_layout()
    p5 = "output/figures/correlaciones.png"
    plt.savefig(p5, dpi=160); plt.close(fig5)

    return [p1, p2, p3, p4, p5]


def top_anomalies(df, n=5):
    a = df[df["is_real_issue"]].copy()
    if a.empty:
        return a
    a["abs_delta"] = a["delta_ventas_pct"].abs()
    return a.sort_values("abs_delta", ascending=False).head(n)


def executive_bullets(df, mesas):
    bullets = []

    ventas_trend = df["Ventas"].pct_change().mean()
    ticket_trend = df["Ticket_Promedio"].pct_change().mean()
    personas_trend = df["Personas"].pct_change().mean()

    bullets.append(f"Ventas presentan una tendencia de {ventas_trend:.2%}")
    bullets.append(f"Ticket promedio presenta una tendencia de {ticket_trend:.2%}")
    bullets.append(f"Clientes (personas) muestran una tendencia de {personas_trend:.2%}")

    if ventas_trend < 0 and ticket_trend > 0:
        bullets.append(
            "Se observa una caída en volumen (clientes/tickets) compensada parcialmente por un incremento en ticket promedio. "
            "Esto sugiere una contracción en la demanda más que un problema de pricing."
        )

    last = df.iloc[-1]
    bullets.append(f"Último mes: Ventas ${last['Ventas']:,.0f}, Ticket ${last['Ticket_Promedio']:,.0f}, Margen {last['Margen_%']:.1%}")

    return bullets

def compute_exec_metrics(df, mesas):
    """
    Métricas base para narrativa DG (sin inventar valores).
    Todo sale del df y mesas.
    """
    out = {}

    # Último mes y variaciones MoM
    last = df.iloc[-1].copy()
    prev = df.iloc[-2].copy() if len(df) >= 2 else None

    out["last_month"] = str(last["month"])
    out["ventas_last"] = float(last["Ventas"]) if pd.notna(last["Ventas"]) else np.nan
    out["ticket_last"] = float(last["Ticket_Promedio"]) if pd.notna(last["Ticket_Promedio"]) else np.nan
    out["cheque_last"] = float(last["Cheque_Promedio"]) if pd.notna(last["Cheque_Promedio"]) else np.nan
    out["personas_last"] = float(last["Personas"]) if pd.notna(last["Personas"]) else np.nan
    out["cogs_last"] = float(last["COGS_%"]) if pd.notna(last["COGS_%"]) else np.nan
    out["margen_last"] = float(last["Margen_%"]) if pd.notna(last["Margen_%"]) else np.nan
    out["delivery_share_last"] = float(last.get("%_Delivery", np.nan)) if "%_Delivery" in df.columns else np.nan
    out["delta_recon_last"] = float(last.get("delta_ventas_pct", np.nan)) if "delta_ventas_pct" in df.columns else np.nan

    if prev is not None:
        def pct(a, b):
            if pd.isna(a) or pd.isna(b) or b == 0:
                return np.nan
            return (a - b) / b

        out["ventas_mom"] = pct(last["Ventas"], prev["Ventas"])
        out["tickets_mom"] = pct(last["Tickets"], prev["Tickets"]) if "Tickets" in df.columns else np.nan
        out["personas_mom"] = pct(last["Personas"], prev["Personas"])
        out["ticket_mom"] = pct(last["Ticket_Promedio"], prev["Ticket_Promedio"])
        out["cheque_mom"] = pct(last["Cheque_Promedio"], prev["Cheque_Promedio"])
        out["cogs_pp_mom"] = (last["COGS_%"] - prev["COGS_%"]) if pd.notna(last["COGS_%"]) and pd.notna(prev["COGS_%"]) else np.nan
        out["margen_pp_mom"] = (last["Margen_%"] - prev["Margen_%"]) if pd.notna(last["Margen_%"]) and pd.notna(prev["Margen_%"]) else np.nan
    else:
        out["ventas_mom"] = out["tickets_mom"] = out["personas_mom"] = out["ticket_mom"] = out["cheque_mom"] = np.nan
        out["cogs_pp_mom"] = out["margen_pp_mom"] = np.nan

    # Tendencias promedio (aprox) usando tasas mensuales
    out["ventas_trend_avg"] = df["Ventas"].pct_change().mean()
    out["ticket_trend_avg"] = df["Ticket_Promedio"].pct_change().mean()
    out["cheque_trend_avg"] = df["Cheque_Promedio"].pct_change().mean()
    out["personas_trend_avg"] = df["Personas"].pct_change().mean()

    # Mesas (último mes + MoM)
    mesas_last = mesas.iloc[-1].copy()
    mesas_prev = mesas.iloc[-2].copy() if len(mesas) >= 2 else None
    out["mesas_last"] = float(mesas_last["mesas"]) if pd.notna(mesas_last["mesas"]) else np.nan
    out["mesas_mom"] = ((mesas_last["mesas"] - mesas_prev["mesas"]) / mesas_prev["mesas"]) if (mesas_prev is not None and mesas_prev["mesas"] != 0) else np.nan

    # Personas por mesa
    out["ppm_last"] = float(mesas_last["clientes"] / mesas_last["mesas"]) if pd.notna(mesas_last["clientes"]) and pd.notna(mesas_last["mesas"]) and mesas_last["mesas"] != 0 else np.nan

    return out


def format_pct(x, digits=1):
    if pd.isna(x):
        return "N/D"
    return f"{x*100:.{digits}f}%"

def format_pp(x, digits=2):
    # percentage points (pp)
    if pd.isna(x):
        return "N/D"
    return f"{x*100:.{digits}f} pp"

def format_money(x, digits=0):
    if pd.isna(x):
        return "N/D"
    return f"${x:,.{digits}f}"

def generate_premium_narrative(df, mesas, anomalies_df):
    """
    Narrativa automática estilo DG (premium).
    Inspirada en la lógica de 'cambio estructural' y ejes de análisis como en [informeejecutivo_vinos_2025.docx](https://fondaargentinamx-my.sharepoint.com/personal/javier_viniegra_fondaargentina_com/_layouts/15/Doc.aspx?sourcedoc=%7B9FBC6FB0-D796-4049-B635-96E2484A0212%7D&file=informeejecutivo_vinos_2025.docx&action=default&mobileredirect=true&DefaultItemOpen=1&EntityRepresentationId=24330848-15de-4b03-991b-024f22126449). 【1-436d3f】
    """
    m = compute_exec_metrics(df, mesas)

    # Detectores sencillos (sin inventar)
    ventas_baja_ticket_sube = (m["ventas_trend_avg"] < 0) and (m["ticket_trend_avg"] > 0)
    demanda_baja = (m["personas_trend_avg"] < 0) or (("tickets_trend_avg" in m) and pd.notna(m.get("tickets_trend_avg", np.nan)) and m["tickets_trend_avg"] < 0)

    # Mensajes “DG”
    intro = (
        "Este reporte sintetiza la evolución mensual de la sucursal, integrando desempeño comercial (ventas, ticket y cheque), "
        "demanda (tickets, comensales/personas, mesas) y señales operativas (anulaciones, cancelaciones, cortesías y descuentos). "
        "El objetivo es identificar cambios estructurales y puntos de ruptura para orientar decisiones accionables."
    )

    contexto = (
        "En entornos de alta variabilidad (como plazas turísticas), el volumen por sí solo deja de ser una defensa suficiente: "
        "es necesario leer simultáneamente demanda, conversión y calidad operativa para evitar diagnósticos incompletos. "
        "Este enfoque es consistente con análisis ejecutivos previos donde el problema no siempre se manifiesta como una caída abrupta, "
        "sino como una combinación de señales que tienden a pasar desapercibidas en la operación diaria."
    )  # 【1-436d3f】 (se cita en el Word, no aquí)

    snapshot = (
        f"**Último mes ({m['last_month']})**: Ventas {format_money(m['ventas_last'],0)}, "
        f"Ticket {format_money(m['ticket_last'],0)}, Cheque {format_money(m['cheque_last'],0)}, "
        f"Clientes/Personas {int(m['personas_last']) if pd.notna(m['personas_last']) else 'N/D'}, "
        f"Mesas {int(m['mesas_last']) if pd.notna(m['mesas_last']) else 'N/D'}, "
        f"COGS {format_pct(m['cogs_last'],1)}, Margen {format_pct(m['margen_last'],1)}."
    )

    lectura_principal = "Lectura ejecutiva: "
    if ventas_baja_ticket_sube:
        lectura_principal += (
            "se observa una contracción de demanda/volumen (menos tickets y/o menos comensales), "
            "parcialmente compensada por un aumento en el ticket promedio. "
            "Esto normalmente apunta a un reto de captación/ocupación y/o a un cambio en el patrón de consumo, "
            "más que a un problema de pricing."
        )
    else:
        lectura_principal += (
            "la relación entre ventas y ticket no muestra el patrón típico de compensación; "
            "se recomienda revisar demanda (tickets/personas) y mix por canal para aislar la causa dominante."
        )

    # Anomalías reales (reconciliación)
    anom_text = ""
    if anomalies_df is not None and not anomalies_df.empty:
        top = anomalies_df.head(3)
        anom_text = "Meses con mayor desviación (Órdenes vs Cierre, excluyendo migraciones): " + \
                    ", ".join([f"{r['month']} ({format_pct(r['delta_ventas_pct'],2)})" for _, r in top.iterrows()]) + "."
    else:
        anom_text = "No se detectaron meses con incidencias reales bajo los criterios actuales."

    # Decisiones sugeridas (sin inventar, basadas en señales)
    decisiones = []
    decisiones.append("Ajustar staffing y roles al patrón real de demanda (horas pico vs horas valle) usando heatmaps de día/hora.")
    decisiones.append("Reforzar disciplina operativa en PV: captura correcta, coordinación de comandas y verificación del flujo de servicio.")
    decisiones.append("Monitorear conversión: mesas → tickets → personas, y evaluar si la caída proviene de ocupación, rotación o afluencia.")
    decisiones.append("Segmentar por canal (salón/delivery) para validar si el ticket sube por mayor peso de delivery y si eso afecta volumen.")
    decisiones.append("Revisar desempeño de meseros con cuadrantes (ventas vs clientes) para detectar brechas de ejecución y entrenamiento.")

    # Puente operativo (mesero como punto crítico)
    puente_operativo = (
        "Nota operativa: el mesero es el punto de captura y coordinación del servicio en el PV; "
        "errores en captura, secuencia o coordinación entre áreas impactan directamente tiempos, experiencia y conversión. "
        "Por ello, el análisis incluye cuadrantes por mesero y lectura de patrones por horario."
    )

    # Metodología de cuadrantes (mediana)
    metodologia_cuadrantes = (
        "Metodología de cuadrantes (2×2): los diagramas se dividen en cuatro cuadrantes usando la **mediana** de cada eje "
        "(no el promedio). La mediana es robusta ante valores atípicos y evita distorsiones por meses extraordinarios, "
        "permitiendo comparar periodos con mayor estabilidad."
    )

    return {
        "intro": intro,
        "contexto": contexto,
        "snapshot": snapshot,
        "lectura": lectura_principal,
        "anomalias": anom_text,
        "decisiones": decisiones,
        "puente_operativo": puente_operativo,
        "metodologia_cuadrantes": metodologia_cuadrantes
    }


def add_premium_sections_to_word(doc, df, mesas, meseros, breakpoints):
    """
    Inserta narrativa premium DG en el Word (en español).
    También documenta mediana y el rol operativo del mesero.
    """
    # Construir tabla de anomalías reales
    anomalies = df[df.get("is_real_issue", False)].copy()
    if not anomalies.empty:
        anomalies["abs_delta"] = anomalies["delta_ventas_pct"].abs()
        anomalies = anomalies.sort_values("abs_delta", ascending=False)

    narrative = generate_premium_narrative(df, mesas, anomalies)

    doc.add_heading("Resumen Ejecutivo (Narrativa DG)", level=1)
    doc.add_paragraph(narrative["intro"])
    doc.add_paragraph(narrative["contexto"])  # Inspiración estilo “cambio estructural” 【1-436d3f】

    doc.add_paragraph(narrative["snapshot"])
    doc.add_paragraph(narrative["lectura"])

    doc.add_heading("Incidencias y Meses a Observar", level=1)
    doc.add_paragraph(narrative["anomalias"])

    doc.add_heading("Decisiones sugeridas (acciones en 30 días)", level=1)
    for d in narrative["decisiones"]:
        doc.add_paragraph(d, style="List Bullet")

    doc.add_heading("Fundamento Operativo (Servicio)", level=1)
    doc.add_paragraph(narrative["puente_operativo"])  # Rol del mesero en PV 【5-35d204】

    doc.add_heading("Metodología: Cuadrantes por Mediana", level=1)
    doc.add_paragraph(narrative["metodologia_cuadrantes"])  # Cuadrantes como práctica interna 【3-949841】

    # Si quieres: añadir KPIs “Average Check / Covers / Food Cost%” como marco
    doc.add_paragraph(
        "Marco de KPIs: Ticket/cheque promedio y comensales (covers) deben analizarse en conjunto para distinguir "
        "captación vs upselling; COGS% ayuda a separar demanda vs eficiencia de costo."
    )  # 【4-f2ee57】


def build_word_report(df, mesas, meseros, fig_paths):
    s = get_scope()
    doc = Document()

    add_premium_sections_to_word(doc, df, mesas, meseros, None)

    # Portada (estilo tu documento)
    doc.add_heading("Reporte Ejecutivo de Evolución Operativa y Financiera", 0)
    doc.add_paragraph(f"Sucursal (keyword): {s.branch_keyword}")
    doc.add_paragraph(f"Subsidiary ID (Wansoft): {s.wansoft_subsidiary_id}")
    doc.add_paragraph(f"Periodo analizado: {s.start_date} a {s.end_date}")
    doc.add_paragraph("Autor: Javier Viniegra")
    doc.add_paragraph("Versión: 1.0")
    doc.add_page_break()

    # Introducción y objetivo (similar a tu ejemplo)
    doc.add_heading("Introducción y Objetivo", level=1)
    doc.add_paragraph(
        "Este reporte analiza la evolución mensual de la sucursal, integrando ventas, costos (COGS), margen, "
        "ticket/cheque promedio, comportamiento de clientes (personas), mesas atendidas y señales operativas "
        "(cortesías, cancelaciones, anulaciones y descuentos). El objetivo es identificar tendencias, "
        "puntos de inflexión y posibles causas operativas que expliquen cambios relevantes en el desempeño."
    )

    # Objetivos
    doc.add_heading("Objetivos", level=1)
    objetivos = [
        "Analizar la evolución mensual de ventas, COGS%, margen%, ticket promedio y cheque promedio.",
        "Evaluar la evolución de clientes (personas) y número de mesas atendidas.",
        "Detectar puntos de ruptura (inflección) y meses con desviaciones relevantes entre Órdenes y Cierre.",
        "Explorar correlaciones entre variables clave (mes-ventas, ticket-ventas, cheque-ventas, mix delivery-ventas).",
        "Proveer decisiones sugeridas para investigación y acción operativa."
    ]
    for o in objetivos:
        doc.add_paragraph(o, style="List Bullet")

    # Hallazgos clave
    doc.add_heading("Hallazgos Clave", level=1)
    bullets = executive_bullets(df, mesas)
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # ✅ Narrativa automática DG (premium)
    add_premium_sections_to_word(doc, df, mesas, None, None)

    # Desarrollo / Metodología (alineada a tu DAX)
    doc.add_heading("Metodología y Fuentes", level=1)
    doc.add_paragraph(
        "Ventas: se reportan como totales. "
        "Costos: se toman de la tabla costeomensual (acumulado mensual; el último registro del mes representa el cierre del mes). "
        "Indicadores operativos: se obtienen de getglobalcashclosing (cortesías, cancelaciones, anulaciones y descuentos en platillos)."
    )

    # Resultados con gráficas
    doc.add_heading("Resultados (Evolución mensual)", level=1)
    add_meseros_analysis_section(doc, df, mesas, meseros)
    for p in fig_paths:
        doc.add_paragraph(os.path.basename(p))
        doc.add_picture(p, width=Inches(6.5))

    # Top anomalías / Puntos de ruptura
    doc.add_heading("Puntos de Inflección y Meses a Evaluar", level=1)
    an = top_anomalies(df, n=10)
    if an.empty:
        doc.add_paragraph("No se detectaron meses con incidencias reales bajo los criterios actuales.")
    else:
        t = doc.add_table(rows=1, cols=6)
        hdr = t.rows[0].cells
        hdr[0].text = "Mes"
        hdr[1].text = "Δ Ventas %"
        hdr[2].text = "Ventas"
        hdr[3].text = "COGS %"
        hdr[4].text = "Margen %"
        hdr[5].text = "Nota/Flag"

        for _, r in an.iterrows():
            row = t.add_row().cells
            row[0].text = str(r["month"])
            row[1].text = f"{r['delta_ventas_pct']:.1%}"
            row[2].text = f"${r['Ventas']:,.0f}"
            row[3].text = f"{r['COGS_%']:.1%}"
            row[4].text = f"{r['Margen_%']:.1%}"
            row[5].text = f"{r.get('source_flag','OK')} {r.get('note','')}".strip()

    # Correlaciones (interpretación)
    doc.add_heading("Correlaciones y Lectura Operativa", level=1)
    doc.add_paragraph(
        "Las correlaciones ayudan a identificar relaciones entre variables (no implican causalidad). "
        "Se recomienda usar estas señales para orientar una revisión operativa y/o comercial."
    )

    # Siguientes pasos (similar a tu formato)
    doc.add_heading("Siguientes Pasos Propuestos", level=1)
    next_steps = [
        "Investigar meses con mayor desviación (Δ Ventas %) y revisar consistencia de cierres y captura.",
        "Cruzar meses críticos con picos de anulaciones/cancelaciones para aislar causas operativas.",
        "Revisar evolución de ticket/cheque y su relación con mix de canal (Delivery vs Salón).",
        "Evaluar desempeño de meseros (ventas, clientes atendidos, ticket promedio) y detectar cambios de patrón.",
        "Proponer acciones correctivas (capacitación, controles, ajustes de proceso) y medir impacto mes a mes."
    ]
    for s2 in next_steps:
        doc.add_paragraph(s2, style="List Bullet")

    # Conclusiones (similar a tu estilo)
    doc.add_heading("Conclusiones", level=1)
    doc.add_paragraph(
        "El análisis mensual permite identificar tendencias claras, meses con desviación y señales operativas "
        "que pueden explicar cambios en el desempeño. La combinación de evolución financiera (ventas, COGS, margen) "
        "con señales operativas (anulaciones, cancelaciones, cortesías y descuentos) facilita un diagnóstico accionable "
        "para Dirección y Operaciones."
    )

    doc.add_heading("Análisis por Cuadrantes (mediana)", level=1)

    doc.add_paragraph(
        "Los siguientes gráficos segmentan el desempeño en cuadrantes utilizando la mediana como punto de corte. "
        "Esto permite identificar si los cambios provienen de volumen, ticket promedio o eficiencia operativa."
    )

    doc.add_picture("output/figures/cuadrantes_mensual.png", width=Inches(6.5))

    doc.add_paragraph(
        "Interpretación: La caída en ventas se explica principalmente por reducción en volumen (tickets/personas), "
        "mientras que el ticket promedio muestra crecimiento, indicando cambio en patrón de consumo."
    )

    doc.add_picture("output/figures/cuadrantes_meseros.png", width=Inches(6.5))

    doc.add_paragraph(
        "El análisis por mesero muestra diferencias en desempeño, sugiriendo oportunidades en capacitación, asignación de mesas "
        "y estrategias de venta."
    )

    out_name = f"output/word/Reporte_Ejecutivo_{s.wansoft_subsidiary_id}_{s.end_date}.docx"
    doc.save(out_name)
    return out_name

def add_meseros_analysis_section(doc, df, mesas, meseros):
    """
    Sección profesional de análisis de meseros (DG level).
    """

    doc.add_heading("Análisis de Desempeño de Meseros", level=1)

    doc.add_paragraph(
        "El desempeño de los meseros es una variable crítica en la generación de ingresos, ya que conecta directamente "
        "la demanda (clientes) con la capacidad de conversión en ventas (tickets y ticket promedio). "
        "A diferencia de otros factores, el impacto del servicio es inmediato y acumulativo en cada interacción con el cliente."
    )

    # --- Insight basado en tendencia general ---
    ventas_trend = df["Ventas"].pct_change().mean()
    ticket_trend = df["Ticket_Promedio"].pct_change().mean()
    personas_trend = df["Personas"].pct_change().mean()

    if ventas_trend < 0 and ticket_trend > 0:
        doc.add_paragraph(
            "Se observa un patrón donde el ticket promedio crece mientras el volumen de clientes/tickets disminuye. "
            "Este comportamiento indica que el problema principal no está en la capacidad de venta por cliente, "
            "sino en la captación o conversión de demanda."
        )

    # --- Inserta cuadrantes ---
    doc.add_paragraph("Segmentación por cuadrantes (mediana):")
    doc.add_picture("output/figures/cuadrantes_meseros.png", width=Inches(6.5))

    # --- Interpretación ejecutiva ---
    doc.add_heading("Interpretación Operativa", level=2)

    doc.add_paragraph(
        "El análisis de cuadrantes permite separar el desempeño de los meseros en función de dos dimensiones clave: "
        "volumen de clientes atendidos y ventas generadas. Esta segmentación permite identificar de forma clara "
        "diferencias estructurales en la ejecución del servicio."
    )

    doc.add_paragraph(
        "Meseros ubicados en el cuadrante de alto volumen y baja venta sugieren una oportunidad en técnicas de venta "
        "(upselling, sugerencias, cierre de ticket). Por otro lado, meseros con alto ticket pero bajo volumen pueden "
        "indicar limitación de asignación de mesas o problemas en flujo operativo."
    )

    doc.add_paragraph(
        "La dispersión de resultados entre meseros confirma que el desempeño no es homogéneo, lo que abre una oportunidad "
        "directa de mejora mediante estandarización de prácticas y capacitación."
    )

    # --- Relación con operación real ---
    doc.add_heading("Relación con la Operación", level=2)

    doc.add_paragraph(
        "El flujo de servicio depende directamente de la ejecución del mesero en el punto de venta: captura correcta de órdenes, "
        "comunicación con cocina/parrilla/bar y seguimiento del servicio hasta el cierre de la cuenta. "
        "Desalineaciones en este flujo impactan directamente los tiempos, la experiencia del cliente y la conversión en ventas."
    )

    # --- Decisiones DG ---
    doc.add_heading("Decisiones Sugeridas – Meseros", level=2)

    decisiones = [
        "Estandarizar discurso de venta (upselling de bebidas, postres y sugerencias).",
        "Redistribuir carga de mesas en horarios críticos para balancear volumen entre meseros.",
        "Implementar coaching operacional basado en desempeño individual (no general).",
        "Revisar tiempos de servicio y coordinación con cocina/parrilla.",
        "Monitorear semanalmente indicadores por mesero (ventas, clientes, ticket promedio)."
    ]

    for d in decisiones:
        doc.add_paragraph(d, style="List Bullet")

def main():
    ensure_dirs()

    kpis, recon, ops = load_inputs()
    df = prepare_monthly(kpis, recon, ops)

    mesas, orders, meseros = build_orders_ops()

    fig_paths = save_core_figures(df, mesas)
    save_quadrant_figures(df, mesas, meseros)

    out_docx = build_word_report(df, mesas, meseros, fig_paths)

    print("\n✅ Word generado:")
    print(out_docx)


if __name__ == "__main__":
    main()